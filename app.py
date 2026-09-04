import os
import json
import time
import requests
import pandas as pd
import numpy as np
import io
import pdfplumber
import docx
from bs4 import BeautifulSoup
from dotenv import load_dotenv
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity
from google import genai
from google.genai import types
from google.genai.errors import ServerError
import streamlit as st

# Direct imports from your existing helper modules
from structured_output_helper import MatchReport, JobEvaluation
from skill_recommender import UpskillingRoadmap, get_course_url

# Load environment variables
load_dotenv(override=True)

# -----------------------------------------------------------------------------
# Streamlit Page Config
# -----------------------------------------------------------------------------
st.set_page_config(
    page_title="AI Job Match & Skill Gap Recommender",
    page_icon="🎯",
    layout="wide"
)

# -----------------------------------------------------------------------------
# Core Pipeline Functions (Mirrors ResumeJoogle.ipynb)
# -----------------------------------------------------------------------------
def clean_html(raw_html: str) -> str:
    """Strips HTML tags and normalizes whitespace."""
    if not raw_html:
        return ""
    soup = BeautifulSoup(raw_html, "html.parser")
    text = soup.get_text(separator=" ")
    return " ".join(text.split())

@st.cache_data(show_spinner=False)
def fetch_jobs_pipeline(api_key, keywords="AI ML Engineer", location="India", total_pages=3):
    url = f"https://jooble.org/api/{api_key}"
    all_jobs = []

    for page in range(1, total_pages + 1):
        payload = {
            "keywords": keywords,
            "location": location,
            "page": str(page)
        }
        response = requests.post(url, json=payload)
        if response.status_code == 200:
            jobs = response.json().get("jobs", [])
            all_jobs.extend(jobs)
        else:
            break
        time.sleep(0.5)

    df = pd.DataFrame(all_jobs)
    if not df.empty:
        df = df[["id", "title", "company", "location", "salary", "link", "snippet"]]
        df["clean_description"] = df["snippet"].apply(clean_html)
        df = df.drop_duplicates(subset=["id"]).reset_index(drop=True)

    return all_jobs, df

def parse_resume(file_input) -> str:
    """Extracts text from PDF, DOCX, or TXT files."""
    file_name = file_input.name if hasattr(file_input, "name") else str(file_input)
    ext = file_name.split(".")[-1].lower()
    extracted_text = []

    if ext == "pdf":
        with pdfplumber.open(file_input) as pdf:
            for page in pdf.pages:
                text = page.extract_text()
                if text:
                    extracted_text.append(text)
    elif ext == "docx":
        if hasattr(file_input, "read"):
            doc = docx.Document(io.BytesIO(file_input.read()))
        else:
            doc = docx.Document(file_input)
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                extracted_text.append(paragraph.text)
    elif ext == "txt":
        if hasattr(file_input, "read"):
            return file_input.read().decode("utf-8")
        with open(file_input, "r", encoding="utf-8") as f:
            return f.read()
    else:
        raise ValueError(f"Unsupported file format: .{ext}")

    full_text = "\n".join(extracted_text)
    return " ".join(full_text.split())

def rank_jobs_tfidf(resume_text: str, df: pd.DataFrame, top_n: int = 10) -> pd.DataFrame:
    """Ranks job listings against a resume string using TF-IDF and Cosine Similarity."""
    corpus = [resume_text] + df["clean_description"].tolist()
    vectorizer = TfidfVectorizer(stop_words="english", ngram_range=(1, 2))
    tfidf_matrix = vectorizer.fit_transform(corpus)
    similarity_scores = cosine_similarity(tfidf_matrix[0:1], tfidf_matrix[1:]).flatten()

    df_ranked = df.copy()
    df_ranked["tfidf_score"] = similarity_scores
    df_ranked["match_score_%"] = (df_ranked["tfidf_score"] * 100).round(1)
    return df_ranked.sort_values(by="tfidf_score", ascending=False).head(top_n).reset_index(drop=True)

def evaluate_matches_with_gemini(resume_text: str, top_10_df: pd.DataFrame, client: genai.Client) -> MatchReport:
    """Passes top TF-IDF candidate jobs and resume text into Gemini 3.6 Flash."""
    jobs_payload = top_10_df[["id", "title", "company", "clean_description"]].to_dict(orient="records")

    prompt = f"""
    You are an expert AI Technical Recruiter. Evaluate the candidate's resume against the 
    provided shortlist of candidate jobs (pre-filtered by relevance).

    RESUME TEXT:
    ---
    {resume_text[:4000]}
    ---

    CANDIDATE JOBS LIST (JSON):
    ---
    {json.dumps(jobs_payload, indent=2)}
    ---

    Task:
    1. Assess semantic fit (look beyond exact keywords for conceptual overlap like PyTorch vs Deep Learning).
    2. Score fit from 0 to 100.
    3. Identify matching skills and missing technical requirements.
    4. Provide a punchy 2-sentence rationale per role.
    """

    response = client.models.generate_content(
        model="gemini-3.6-flash",
        contents=prompt,
        config=types.GenerateContentConfig(
            response_mime_type="application/json",
            response_schema=MatchReport,
            temperature=0.2,
        ),
    )
    return response.parsed

def generate_upskilling_roadmap(selected_job: pd.Series, client: genai.Client, max_retries: int = 3) -> UpskillingRoadmap:
    """Recommends courses for missing skills using Gemini 3.6 Flash with retry logic."""
    missing_skills = selected_job["missing_skills"]
    if not missing_skills:
        return None

    prompt = f"""
    You are an AI Career Coach. A candidate is applying for the role:
    - Role: {selected_job['title']}
    - Company: {selected_job['company']}
    
    The candidate has identified the following missing skill gaps:
    {missing_skills}

    For EACH missing skill listed above, provide exactly 3 well-known, high-quality, 
    real-world learning resources:
    1. Beginner (foundations & core mental models)
    2. Intermediate (hands-on application & tooling)
    3. Advanced (enterprise scaling, production patterns, or leadership)
    """

    for attempt in range(1, max_retries + 1):
        try:
            response = client.models.generate_content(
                model="gemini-3.6-flash",
                contents=prompt,
                config=types.GenerateContentConfig(
                    response_mime_type="application/json",
                    response_schema=UpskillingRoadmap,
                    temperature=0.2,
                ),
            )
            return response.parsed
        except ServerError as err:
            if attempt < max_retries:
                time.sleep(attempt * 2)
            else:
                raise err

# -----------------------------------------------------------------------------
# Sidebar: Setup
# -----------------------------------------------------------------------------
with st.sidebar:
    st.header("⚙️ Configuration")
    
    env_jooble = os.getenv("api_key", "")
    env_gemini = os.getenv("GEMINI_API_KEY", "")

    jooble_api_key = st.text_input("Jooble API Key", value=env_jooble, type="password")
    gemini_api_key = st.text_input("Gemini API Key", value=env_gemini, type="password")
    
    st.markdown("---")
    st.subheader("Search Parameters")
    search_keywords = st.text_input("Keywords", value="AI ML Engineer")
    search_location = st.text_input("Location", value="India")
    search_pages = st.slider("Pages to Ingest", min_value=1, max_value=4, value=3)

# -----------------------------------------------------------------------------
# Main Application UI
# -----------------------------------------------------------------------------
st.title("🎯 AI Job Match & Skill Gap Engine")
st.caption("End-to-End Pipeline: Jooble API ➔ TF-IDF Filter ➔ Gemini 3.6 Flash Evaluation ➔ Course Roadmap")

uploaded_file = st.file_uploader("Upload Resume (.pdf or .docx)", type=["pdf", "docx"])

if uploaded_file is not None:
    sample_resume_text = parse_resume(uploaded_file)
    st.success(f"Resume parsed ({len(sample_resume_text):,} characters).")
    
    with st.expander("📄 View Parsed Resume Text Preview"):
        st.write(sample_resume_text[:1000] + "...")

    if st.button("🚀 Find & Evaluate Matching Jobs", type="primary"):
        if not jooble_api_key or not gemini_api_key:
            st.error("Please ensure Jooble and Gemini API keys are provided.")
            st.stop()

        client = genai.Client(api_key=gemini_api_key)

        with st.status("Running Pipeline...", expanded=True) as status:
            status.write("📥 Step 1: Fetching jobs from Jooble...")
            _, df_jobs = fetch_jobs_pipeline(
                api_key=jooble_api_key,
                keywords=search_keywords,
                location=search_location,
                total_pages=search_pages
            )
            status.write(f"✅ Ingested {len(df_jobs)} jobs.")

            status.write("📊 Step 2: Running TF-IDF coarse ranking...")
            top_10_jobs = rank_jobs_tfidf(sample_resume_text, df_jobs, top_n=10)
            status.write("✅ Filtered to Top 10 jobs.")

            status.write("🤖 Step 3: Running semantic evaluation via Gemini 3.6 Flash...")
            report: MatchReport = evaluate_matches_with_gemini(sample_resume_text, top_10_jobs, client)
            status.write("✅ Structured evaluation received.")
            status.update(label="Complete!", state="complete", expanded=False)

        evaluated_records = [job.model_dump() for job in report.evaluated_jobs]
        df_results = pd.DataFrame(evaluated_records)
        df_results = df_results.sort_values(by="semantic_fit_score", ascending=False).reset_index(drop=True)

        st.session_state["df_results"] = df_results
        st.session_state["summary"] = report.summary
        st.session_state["client"] = client

# -----------------------------------------------------------------------------
# Results Display
# -----------------------------------------------------------------------------
if "df_results" in st.session_state:
    df_results = st.session_state["df_results"]
    summary = st.session_state["summary"]
    client = st.session_state["client"]

    st.markdown("---")
    st.subheader("💡 Candidate Summary")
    st.info(summary)

    st.subheader("📋 Top Semantic Job Matches")

    for idx, row in df_results.iterrows():
        with st.container(border=True):
            col1, col2 = st.columns([4, 1])
            with col1:
                st.markdown(f"### {row['title']}")
                st.markdown(f"**Company:** `{row['company']}`")
            with col2:
                st.metric("Semantic Fit", f"{row['semantic_fit_score']}%")

            st.write(f"**Rationale:** {row['rationale']}")

            c_left, c_right = st.columns(2)
            with c_left:
                st.markdown("**✅ Matching Skills:**")
                st.write(", ".join([f"`{s}`" for s in row["matching_skills"]]) if row["matching_skills"] else "None")

            with c_right:
                st.markdown("**⚠️ Missing Skills:**")
                st.write(", ".join([f"`{s}`" for s in row["missing_skills"]]) if row["missing_skills"] else "🎉 None")

            if row["missing_skills"]:
                with st.expander(f"📚 Bridge Skills Gap for {row['title']}"):
                    btn_key = f"btn_{idx}"
                    roadmap_key = f"roadmap_{idx}"

                    if st.button("Generate Learning Recommendations", key=btn_key):
                        with st.spinner("Curating courses with Gemini 3.6 Flash..."):
                            roadmap = generate_upskilling_roadmap(row, client)
                            st.session_state[roadmap_key] = roadmap

                    if roadmap_key in st.session_state and st.session_state[roadmap_key] is not None:
                        current_roadmap: UpskillingRoadmap = st.session_state[roadmap_key]
                        for path in current_roadmap.learning_paths:
                            st.markdown(f"#### Skill: `{path.skill_name.upper()}`")
                            for course in path.courses:
                                link = get_course_url(course)
                                st.markdown(f"- **[{course.level}]** [{course.course_name}]({link}) — *{course.platform_or_provider}*")
                                st.caption(f"  ↳ {course.key_takeaway}")
                            st.divider()